#!/usr/bin/env python3
"""
Interactive console RAG chatbot with persistent thread memory.
Supports OpenAI and vLLM providers via factory pattern.
"""

import argparse
import json
import logging
import os
import sys
import time
import uuid
from datetime import datetime, timedelta
from glob import glob
from pathlib import Path
from typing import Dict, List, Optional, Any
import threading

# Add parent directory to path for imports
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from dotenv import load_dotenv
from langchain_core.documents import Document
from langchain_core.chat_history import BaseChatMessageHistory
from langchain_core.messages import AIMessage, BaseMessage, HumanMessage
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.runnables import RunnablePassthrough, RunnableLambda
from langchain_core.runnables.history import RunnableWithMessageHistory
from langchain_openai import ChatOpenAI
from pydantic import BaseModel
from tqdm import tqdm

# Import our RAG package (installed as local dependency)
from rag_workflow import create_rag_workflow, RAGState, create_conversation_summary_prompt

# Import ingestion module components
try:
    from ingestion import ingest_pdf, MemoryStore, DocumentCtx
    from ingestion.config import get_config
    HAS_INGESTION = True
except ImportError:
    HAS_INGESTION = False
    print("Warning: ingestion module not found. Install it as a submodule.")

# Import our memory retriever adapter
from src.adapters import MemoryRetrieverAdapter

# Enhanced PDF support with multiple fallback options
PDF_READERS = {}

# Try PyPDF2 first (most common)
try:
    from PyPDF2 import PdfReader as PyPDF2Reader
    PDF_READERS['pypdf2'] = PyPDF2Reader
except ImportError:
    pass

# Try pypdf (newer version)
try:
    from pypdf import PdfReader as PyPdfReader
    PDF_READERS['pypdf'] = PyPdfReader
except ImportError:
    pass

# Try pdfplumber (better for complex layouts)
try:
    import pdfplumber
    PDF_READERS['pdfplumber'] = pdfplumber
except ImportError:
    pass

# Try pymupdf (excellent performance)
try:
    import fitz  # PyMuPDF
    PDF_READERS['pymupdf'] = fitz
except ImportError:
    pass

HAS_PDF = bool(PDF_READERS)

# Set up logging
logging.basicConfig(
    level=logging.WARNING,  # Default to less verbose
    format='%(message)s',   # Simpler format
    handlers=[
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger("RAG")


class ThreadMetadata(BaseModel):
    """Metadata for thread management."""
    thread_id: str
    user_id: str
    created_at: datetime
    last_accessed: datetime
    message_count: int = 0
    title: Optional[str] = None
    tags: List[str] = []
    summary: Optional[str] = None  # Conversation summary
    last_summarized_count: int = 0  # Last message count when summary was updated


class EnhancedChatMessageHistory(BaseChatMessageHistory, BaseModel):
    """Enhanced chat message history with metadata, size limits, and summarization."""
    
    messages: List[BaseMessage] = []
    metadata: ThreadMetadata
    max_messages: int = 20  # Reduced limit since we use summarization
    summarize_threshold: int = 12  # Start summarizing after this many messages
    llm: Optional[Any] = None  # LLM for summarization
    
    def add_message(self, message: BaseMessage) -> None:
        self.messages.append(message)
        self.metadata.message_count += 1
        self.metadata.last_accessed = datetime.now()
        
        # Check if we need to summarize and truncate
        if len(self.messages) > self.summarize_threshold:
            self._update_summary_if_needed()
            
        # Enforce message limit (keep recent messages)
        if len(self.messages) > self.max_messages:
            removed_count = len(self.messages) - self.max_messages
            self.messages = self.messages[-self.max_messages:]
            logger.debug(f"Removed {removed_count} old messages from thread {self.metadata.thread_id}")
    
    def _update_summary_if_needed(self) -> None:
        """Update conversation summary when needed."""
        messages_since_summary = len(self.messages) - self.metadata.last_summarized_count
        
        # Update summary every 8 new messages
        if messages_since_summary >= 8:
            try:
                # Get messages to summarize (exclude very recent ones)
                messages_to_summarize = self.messages[:-4] if len(self.messages) > 4 else self.messages[:-2]
                
                if messages_to_summarize and self.llm:
                    new_summary = create_conversation_summary(
                        messages_to_summarize,
                        self.llm,
                        self.metadata.summary
                    )
                    
                    if new_summary:
                        self.metadata.summary = new_summary
                        self.metadata.last_summarized_count = len(self.messages) - 4
                        logger.info(f"Updated conversation summary for thread {self.metadata.thread_id}")
                        
            except Exception as e:
                logger.error(f"Failed to update summary: {e}")
    
    def clear(self) -> None:
        self.messages.clear()
        self.metadata.message_count = 0
        self.metadata.last_accessed = datetime.now()
        self.metadata.summary = None
        self.metadata.last_summarized_count = 0
    
    def get_summary(self) -> Dict[str, Any]:
        """Get thread summary for management commands."""
        return {
            "thread_id": self.metadata.thread_id,
            "user_id": self.metadata.user_id,
            "title": self.metadata.title or f"Thread {self.metadata.thread_id[:8]}",
            "message_count": len(self.messages),
            "created_at": self.metadata.created_at.isoformat(),
            "last_accessed": self.metadata.last_accessed.isoformat(),
            "tags": self.metadata.tags
        }


class ThreadSafeMemoryManager:
    """Thread-safe memory manager for chat sessions."""
    
    def __init__(self, state_dir: str = ".state", max_threads_per_user: int = 10, llm=None):
        self.state_dir = Path(state_dir)
        self.state_dir.mkdir(exist_ok=True)
        self.max_threads_per_user = max_threads_per_user
        self.llm = llm  # LLM for summarization
        
        # Thread-safe caches
        self._history_cache: Dict[str, EnhancedChatMessageHistory] = {}
        self._user_threads: Dict[str, List[str]] = {}  # user_id -> [thread_ids]
        self._lock = threading.RLock()
        
        # Load existing thread mappings
        self._load_user_mappings()
    
    def _get_thread_file_path(self, thread_id: str) -> Path:
        """Get file path for thread storage."""
        return self.state_dir / f"thread_{thread_id}.json"
    
    def _get_user_mapping_file(self) -> Path:
        """Get file path for user->threads mapping."""
        return self.state_dir / "user_mappings.json"
    
    def _load_user_mappings(self) -> None:
        """Load user->thread mappings from disk."""
        mapping_file = self._get_user_mapping_file()
        if mapping_file.exists():
            try:
                with open(mapping_file, "r") as f:
                    self._user_threads = json.load(f)
                logger.debug(f"Loaded user mappings for {len(self._user_threads)} users")
            except Exception as e:
                logger.warning(f"Could not load user mappings: {e}")
                self._user_threads = {}
    
    def _save_user_mappings(self) -> None:
        """Save user->thread mappings to disk."""
        mapping_file = self._get_user_mapping_file()
        try:
            with open(mapping_file, "w") as f:
                json.dump(self._user_threads, f, indent=2)
        except Exception as e:
            logger.warning(f"Could not save user mappings: {e}")
    
    def get_or_create_thread(self, thread_id: str, user_id: str) -> EnhancedChatMessageHistory:
        """Get or create a thread for a user."""
        with self._lock:
            if thread_id in self._history_cache:
                # Update access time
                self._history_cache[thread_id].metadata.last_accessed = datetime.now()
                return self._history_cache[thread_id]
            
            # Try to load from disk
            thread_file = self._get_thread_file_path(thread_id)
            if thread_file.exists():
                history = self._load_thread_from_disk(thread_id, user_id)
                if history:
                    self._history_cache[thread_id] = history
                    return history
            
            # Create new thread
            return self._create_new_thread(thread_id, user_id)
    
    def _load_thread_from_disk(self, thread_id: str, user_id: str) -> Optional[EnhancedChatMessageHistory]:
        """Load thread from disk."""
        thread_file = self._get_thread_file_path(thread_id)
        try:
            with open(thread_file, "r") as f:
                data = json.load(f)
            
            # Create metadata
            metadata = ThreadMetadata(
                thread_id=thread_id,
                user_id=user_id,
                created_at=datetime.fromisoformat(data["metadata"]["created_at"]),
                last_accessed=datetime.now(),
                message_count=data["metadata"]["message_count"],
                title=data["metadata"].get("title"),
                tags=data["metadata"].get("tags", []),
                summary=data["metadata"].get("summary"),
                last_summarized_count=data["metadata"].get("last_summarized_count", 0)
            )
            
            # Create history
            history = EnhancedChatMessageHistory(metadata=metadata, llm=self.llm)
            
            # Load messages
            for msg_data in data.get("messages", []):
                if msg_data["type"] == "human":
                    history.add_message(HumanMessage(content=msg_data["content"]))
                elif msg_data["type"] == "ai":
                    history.add_message(AIMessage(content=msg_data["content"]))
            
            logger.info(f"Loaded thread {thread_id} with {len(history.messages)} messages")
            return history
            
        except Exception as e:
            logger.warning(f"Could not load thread {thread_id}: {e}")
            return None
    
    def _create_new_thread(self, thread_id: str, user_id: str) -> EnhancedChatMessageHistory:
        """Create a new thread."""
        metadata = ThreadMetadata(
            thread_id=thread_id,
            user_id=user_id,
            created_at=datetime.now(),
            last_accessed=datetime.now()
        )
        
        history = EnhancedChatMessageHistory(metadata=metadata, llm=self.llm)
        self._history_cache[thread_id] = history
        
        # Update user mappings
        if user_id not in self._user_threads:
            self._user_threads[user_id] = []
        
        if thread_id not in self._user_threads[user_id]:
            self._user_threads[user_id].append(thread_id)
            
            # Enforce thread limit per user
            if len(self._user_threads[user_id]) > self.max_threads_per_user:
                # Remove oldest thread
                oldest_thread = self._user_threads[user_id].pop(0)
                self._delete_thread(oldest_thread)
                logger.info(f"Removed oldest thread {oldest_thread} for user {user_id}")
        
        self._save_user_mappings()
        logger.info(f"Created new thread {thread_id} for user {user_id}")
        return history
    
    def save_thread(self, thread_id: str) -> None:
        """Save thread to disk."""
        with self._lock:
            if thread_id not in self._history_cache:
                return
            
            history = self._history_cache[thread_id]
            thread_file = self._get_thread_file_path(thread_id)
            
            # Prepare data for saving
            messages_data = []
            for msg in history.messages:
                if isinstance(msg, HumanMessage):
                    messages_data.append({"type": "human", "content": msg.content})
                elif isinstance(msg, AIMessage):
                    messages_data.append({"type": "ai", "content": msg.content})
            
            data = {
                "metadata": {
                    "thread_id": history.metadata.thread_id,
                    "user_id": history.metadata.user_id,
                    "created_at": history.metadata.created_at.isoformat(),
                    "last_accessed": history.metadata.last_accessed.isoformat(),
                    "message_count": history.metadata.message_count,
                    "title": history.metadata.title,
                    "tags": history.metadata.tags,
                    "summary": history.metadata.summary,
                    "last_summarized_count": history.metadata.last_summarized_count
                },
                "messages": messages_data
            }
            
            try:
                with open(thread_file, "w") as f:
                    json.dump(data, f, indent=2)
            except Exception as e:
                logger.warning(f"Could not save thread {thread_id}: {e}")
    
    def _delete_thread(self, thread_id: str) -> None:
        """Delete a thread from memory and disk."""
        # Remove from cache
        if thread_id in self._history_cache:
            del self._history_cache[thread_id]
        
        # Remove from disk
        thread_file = self._get_thread_file_path(thread_id)
        if thread_file.exists():
            try:
                thread_file.unlink()
            except Exception as e:
                logger.warning(f"Could not delete thread file {thread_id}: {e}")
    
    def clear_thread(self, thread_id: str) -> None:
        """Clear thread messages but keep metadata."""
        with self._lock:
            if thread_id in self._history_cache:
                self._history_cache[thread_id].clear()
                self.save_thread(thread_id)
    
    def get_user_threads(self, user_id: str) -> List[Dict[str, Any]]:
        """Get all threads for a user."""
        with self._lock:
            thread_ids = self._user_threads.get(user_id, [])
            threads = []
            
            for tid in thread_ids:
                if tid in self._history_cache:
                    threads.append(self._history_cache[tid].get_summary())
                else:
                    # Try to load basic info from disk
                    thread_file = self._get_thread_file_path(tid)
                    if thread_file.exists():
                        try:
                            with open(thread_file, "r") as f:
                                data = json.load(f)
                            threads.append({
                                "thread_id": tid,
                                "user_id": user_id,
                                "title": data["metadata"].get("title", f"Thread {tid[:8]}"),
                                "message_count": data["metadata"]["message_count"],
                                "created_at": data["metadata"]["created_at"],
                                "last_accessed": data["metadata"]["last_accessed"],
                                "tags": data["metadata"].get("tags", [])
                            })
                        except Exception:
                            pass
            
            # Sort by last accessed time (most recent first)
            threads.sort(key=lambda x: x.get("last_accessed", ""), reverse=True)
            return threads
    
    def cleanup_old_threads(self, days: int = 30) -> int:
        """Clean up threads older than specified days."""
        cutoff_date = datetime.now() - timedelta(days=days)
        cleaned_count = 0
        
        with self._lock:
            to_remove = []
            
            for user_id, thread_ids in self._user_threads.items():
                for thread_id in thread_ids[:]:  # Copy list to avoid modification during iteration
                    thread_file = self._get_thread_file_path(thread_id)
                    if thread_file.exists():
                        try:
                            with open(thread_file, "r") as f:
                                data = json.load(f)
                            last_accessed = datetime.fromisoformat(data["metadata"]["last_accessed"])
                            
                            if last_accessed < cutoff_date:
                                to_remove.append((user_id, thread_id))
                        except Exception:
                            # If we can't read the file, consider it for removal
                            to_remove.append((user_id, thread_id))
            
            # Remove old threads
            for user_id, thread_id in to_remove:
                if thread_id in self._user_threads[user_id]:
                    self._user_threads[user_id].remove(thread_id)
                self._delete_thread(thread_id)
                cleaned_count += 1
            
            if cleaned_count > 0:
                self._save_user_mappings()
                logger.info(f"Cleaned up {cleaned_count} old threads")
        
        return cleaned_count


# Global memory manager
_memory_manager: Optional[ThreadSafeMemoryManager] = None

# Note: global_llm and global_retriever are no longer needed 
# as they're handled by the RAG workflow package


# RAGState is now imported from the rag package


def get_session_history(thread_id: str, user_id: str = "default") -> BaseChatMessageHistory:
    """Get or create chat history for a session using enhanced memory manager."""
    global _memory_manager
    if _memory_manager is None:
        # Fallback initialization without LLM (summarization won't work)
        _memory_manager = ThreadSafeMemoryManager()
    
    return _memory_manager.get_or_create_thread(thread_id, user_id)


def save_session_history(thread_id: str) -> None:
    """Persist chat history to disk using enhanced memory manager."""
    global _memory_manager
    if _memory_manager is not None:
        _memory_manager.save_thread(thread_id)


def clear_session_history(thread_id: str) -> None:
    """Clear thread messages using enhanced memory manager."""
    global _memory_manager
    if _memory_manager is not None:
        _memory_manager.clear_thread(thread_id)
        print(f"Cleared thread memory: {thread_id}")


def handle_management_commands(command: str, current_user: str, current_thread: str) -> tuple[bool, Optional[str], Optional[str]]:
    """Handle thread/user management commands.
    
    Returns: (handled, new_user_id, new_thread_id)
    """
    global _memory_manager
    if _memory_manager is None:
        return False, None, None
    
    command = command.lower().strip()
    
    if command == "/threads":
        # List threads for current user
        threads = _memory_manager.get_user_threads(current_user)
        if not threads:
            print(f"No threads found for user '{current_user}'")
        else:
            print(f"\n📋 Threads for user '{current_user}':")
            print("-" * 60)
            for i, thread in enumerate(threads, 1):
                status = "📍 CURRENT" if thread["thread_id"] == current_thread else ""
                print(f"{i:2d}. {thread['title'][:40]:<40} {status}")
                print(f"     ID: {thread['thread_id']}")
                print(f"     Messages: {thread['message_count']:3d} | Last: {thread['last_accessed'][:16]}")
                if thread.get('tags'):
                    print(f"     Tags: {', '.join(thread['tags'])}")
                print()
        return True, None, None
    
    elif command == "/users":
        # List all users with thread counts
        if not _memory_manager._user_threads:
            print("No users found")
        else:
            print("\n👥 All Users:")
            print("-" * 40)
            for user_id, thread_ids in _memory_manager._user_threads.items():
                status = "📍 CURRENT" if user_id == current_user else ""
                print(f"• {user_id:<20} ({len(thread_ids)} threads) {status}")
        return True, None, None
    
    elif command.startswith("/switch"):
        # Switch thread or user
        parts = command.split()
        if len(parts) < 2:
            print("Usage: /switch <thread_id> or /switch user <user_id> [thread_id]")
            return True, None, None
        
        if parts[1] == "user" and len(parts) >= 3:
            # Switch user and optionally thread
            new_user = parts[2]
            new_thread = parts[3] if len(parts) > 3 else str(uuid.uuid4())[:8]
            print(f"Switched to user '{new_user}', thread '{new_thread}'")
            return True, new_user, new_thread
        else:
            # Switch thread only
            new_thread = parts[1]
            print(f"Switched to thread '{new_thread}'")
            return True, None, new_thread
    
    elif command.startswith("/new"):
        # Create new thread
        parts = command.split()
        if len(parts) > 1:
            # Use provided thread name/id
            new_thread = parts[1]
        else:
            # Generate new thread ID
            new_thread = str(uuid.uuid4())[:8]
        print(f"Created new thread '{new_thread}'")
        return True, None, new_thread
    
    elif command.startswith("/title"):
        # Set thread title
        parts = command.split(maxsplit=1)
        if len(parts) < 2:
            print("Usage: /title <thread title>")
            return True, None, None
        
        title = parts[1]
        if current_thread in _memory_manager._history_cache:
            _memory_manager._history_cache[current_thread].metadata.title = title
            _memory_manager.save_thread(current_thread)
            print(f"Set thread title to: {title}")
        else:
            print("Current thread not found in cache")
        return True, None, None
    
    elif command.startswith("/tag"):
        # Add tag to thread
        parts = command.split()
        if len(parts) < 2:
            print("Usage: /tag <tag_name>")
            return True, None, None
        
        tag = parts[1]
        if current_thread in _memory_manager._history_cache:
            history = _memory_manager._history_cache[current_thread]
            if tag not in history.metadata.tags:
                history.metadata.tags.append(tag)
                _memory_manager.save_thread(current_thread)
                print(f"Added tag: {tag}")
            else:
                print(f"Tag '{tag}' already exists")
        else:
            print("Current thread not found in cache")
        return True, None, None
    
    elif command == "/cleanup":
        # Clean up old threads
        cleaned = _memory_manager.cleanup_old_threads(days=30)
        print(f"Cleaned up {cleaned} old threads (>30 days)")
        return True, None, None
    
    elif command == "/info":
        # Show current session info
        print(f"\n📊 Current Session Info:")
        print(f"User ID: {current_user}")
        print(f"Thread ID: {current_thread}")
        
        if current_thread in _memory_manager._history_cache:
            history = _memory_manager._history_cache[current_thread]
            metadata = history.metadata
            print(f"Thread Title: {metadata.title or 'Untitled'}")
            print(f"Messages: {len(history.messages)}")
            print(f"Created: {metadata.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
            print(f"Last Access: {metadata.last_accessed.strftime('%Y-%m-%d %H:%M:%S')}")
            if metadata.tags:
                print(f"Tags: {', '.join(metadata.tags)}")
        
        return True, None, None
    
    return False, None, None


def create_conversation_summary(messages: List[BaseMessage], llm, current_summary: str = None) -> str:
    """Create or update conversation summary using LLM."""
    logger.debug("=" * 70)
    logger.debug("📝 CONVERSATION SUMMARIZATION")
    logger.debug("=" * 70)
    
    if not messages:
        logger.debug("No messages to summarize")
        logger.debug("=" * 70)
        return current_summary or ""
    
    logger.info(f"📝 Summarizing {len(messages)} messages")
    
    # Use the utility function from RAG package to create the prompt
    summary_prompt = create_conversation_summary_prompt(messages, current_summary)
    
    logger.debug(f"Summary prompt length: {len(summary_prompt)} chars")
    logger.debug("Full summarization prompt being sent to vLLM:")
    logger.debug("-" * 50)
    logger.debug(summary_prompt)
    logger.debug("-" * 50)
    
    try:
        if llm:
            summary = llm.invoke(summary_prompt).content.strip()
            logger.info(f"📝 Created summary: {summary}")
            logger.debug(f"Created summary length: {len(summary)} chars")
            logger.debug("=" * 70)
            return summary
        else:
            logger.warning("No LLM available for summarization")
            logger.debug("=" * 70)
            return current_summary or ""
    except Exception as e:
        logger.error(f"Failed to create summary: {e}")
        logger.debug("=" * 70)
        return current_summary or ""


def get_contextual_history(history: "EnhancedChatMessageHistory", recent_count: int = 6) -> str:
    """Get conversation context: summary + recent messages."""
    context_parts = []
    
    # Add summary if available
    if history.metadata.summary:
        context_parts.append(f"Previous conversation summary: {history.metadata.summary}")
    
    # Add recent messages
    if len(history.messages) > 0:
        recent_messages = history.messages[-recent_count:] if len(history.messages) > recent_count else history.messages
        
        if recent_messages:
            context_parts.append("Recent conversation:")
            for msg in recent_messages:
                if isinstance(msg, HumanMessage):
                    context_parts.append(f"Employee: {msg.content}")
                elif isinstance(msg, AIMessage):
                    context_parts.append(f"Assistant: {msg.content}")
    
    return "\n".join(context_parts) if context_parts else ""


def create_llm_provider(provider: str, model: str) -> ChatOpenAI:
    """Factory function to create LLM provider based on configuration."""
    if provider == "openai":
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise ValueError("OPENAI_API_KEY is required for OpenAI provider")
        
        return ChatOpenAI(
            model=model,
            api_key=api_key,
            temperature=0
        )
    
    elif provider == "vllm":
        base_url = os.getenv("VLLM_BASE_URL")
        if not base_url:
            raise ValueError("VLLM_BASE_URL is required for vLLM provider")
        
        api_key = os.getenv("VLLM_API_KEY", "sk-local")  # Default fallback
        vllm_model = os.getenv("VLLM_MODEL", model)  # Use VLLM_MODEL if set, fallback to MODEL
        
        return ChatOpenAI(
            model=vllm_model,
            base_url=base_url,
            api_key=api_key,
            temperature=0
        )
    
    else:
        raise ValueError(f"Unsupported provider: {provider}. Use 'openai' or 'vllm'")


def extract_pdf_content(file_path: str) -> str:
    """Extract text from PDF using multiple fallback methods."""
    content = ""
    errors = []
    
    # Method 1: Try PyPDF2/pypdf (fastest, most common)
    for reader_name in ['pypdf2', 'pypdf']:
        if reader_name in PDF_READERS:
            try:
                logger.debug(f"Trying {reader_name} for {file_path}")
                with open(file_path, 'rb') as f:
                    reader = PDF_READERS[reader_name](f)
                    content = "\n".join(page.extract_text() for page in reader.pages)
                if content.strip():
                    logger.info(f"Successfully extracted PDF content using {reader_name}")
                    return content
            except Exception as e:
                errors.append(f"{reader_name}: {e}")
                continue
    
    # Method 2: Try pdfplumber (better for complex layouts, tables)
    if 'pdfplumber' in PDF_READERS:
        try:
            logger.debug(f"Trying pdfplumber for {file_path}")
            with PDF_READERS['pdfplumber'].open(file_path) as pdf:
                pages = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        pages.append(text)
                content = "\n".join(pages)
            if content.strip():
                logger.info("Successfully extracted PDF content using pdfplumber")
                return content
        except Exception as e:
            errors.append(f"pdfplumber: {e}")
    
    # Method 3: Try PyMuPDF (excellent performance and accuracy)
    if 'pymupdf' in PDF_READERS:
        try:
            logger.debug(f"Trying PyMuPDF for {file_path}")
            doc = PDF_READERS['pymupdf'].open(file_path)
            pages = []
            for page_num in range(doc.page_count):
                page = doc[page_num]
                text = page.get_text()
                if text.strip():
                    pages.append(text)
            content = "\n".join(pages)
            doc.close()
            if content.strip():
                logger.info("Successfully extracted PDF content using PyMuPDF")
                return content
        except Exception as e:
            errors.append(f"pymupdf: {e}")
    
    # If all methods failed, log the errors
    if errors:
        logger.warning(f"All PDF extraction methods failed for {file_path}:")
        for error in errors:
            logger.warning(f"  - {error}")
    
    return content


def load_and_ingest_documents(docs_path: str, memory_store: "MemoryStore", 
                              ctx: "DocumentCtx", chunk_size: int = None, 
                              overlap: int = None) -> int:
    """
    Load documents from directory and ingest them into MemoryStore.
    
    Args:
        docs_path: Path to documents directory
        memory_store: MemoryStore instance to ingest into
        ctx: DocumentCtx with tenant/document context
        chunk_size: Target chunk size (optional, uses config default if None)
        overlap: Overlap between chunks (optional, uses config default if None)
        
    Returns:
        Total number of chunks ingested
    """
    if not HAS_INGESTION:
        raise ImportError("Ingestion module not found. Cannot load documents.")
    
    # Find all PDF files
    pdf_files = glob(os.path.join(docs_path, "**", "*.pdf"), recursive=True)
    
    if not pdf_files:
        print(f"No PDF documents found in {docs_path}")
        return 0
    
    print(f"Loading {len(pdf_files)} PDF files...")
    
    total_chunks = 0
    
    for file_path in tqdm(pdf_files):
        try:
            file_size = os.path.getsize(file_path)
            logger.debug(f"Processing {file_path} ({file_size} bytes)")
            
            # Create a unique document_id based on the file path
            file_name = os.path.basename(file_path)
            doc_id = file_name.replace('.pdf', '').replace(' ', '_')
            
            # Update context for this specific document
            doc_ctx = DocumentCtx(
                tenant_id=ctx.tenant_id,
                owner_user_id=ctx.owner_user_id,
                document_id=doc_id,
                visibility=ctx.visibility,
                embedding_version=ctx.embedding_version
            )
            
            # Read PDF bytes
            with open(file_path, 'rb') as f:
                pdf_bytes = f.read()
            
            # Ingest the PDF using the ingestion module
            # Build kwargs for optional parameters
            ingest_kwargs = {
                'ctx': doc_ctx,
                'filename': file_name,
                'raw_pdf_bytes': pdf_bytes,
                'sink': memory_store,
            }
            if chunk_size is not None:
                ingest_kwargs['chunk_size'] = chunk_size
            if overlap is not None:
                ingest_kwargs['overlap'] = overlap
            
            result = ingest_pdf(**ingest_kwargs)
            
            chunks_saved = result['chunks_saved']
            total_chunks += chunks_saved
            
            logger.info(f"✅ Ingested {file_name}: {chunks_saved} chunks")
            logger.debug(f"   Model: {result.get('embedding_model', 'N/A')} (dim={result.get('embedding_dimension', 'N/A')})")
            logger.debug(f"   Content hash: {result['content_hash']}")
            if 'stats' in result:
                logger.debug(f"   Chunking method: {result['stats'].get('chunking_method', 'N/A')}")
            
        except Exception as e:
            print(f"Warning: Could not ingest {file_path}: {e}")
            logger.debug(f"Error details for {file_path}: {e}", exc_info=True)
    
    print(f"✅ Ingested {total_chunks} total chunks from {len(pdf_files)} documents")
    return total_chunks


def load_documents_legacy(docs_path: str) -> List[Document]:
    """
    Legacy document loader (kept for backward compatibility with non-PDF files).
    This is only used if you need to support non-PDF formats.
    """
    documents = []
    
    # Find all supported file types (excluding PDFs as they use the new ingestion)
    patterns = [
        os.path.join(docs_path, "**", "*.md"),
        os.path.join(docs_path, "**", "*.txt"),
        os.path.join(docs_path, "**", "*.docx"),
    ]
    
    files = []
    for pattern in patterns:
        files.extend(glob(pattern, recursive=True))
    
    if not files:
        return documents
    
    print(f"Loading {len(files)} non-PDF files (legacy mode)...")
    
    for file_path in tqdm(files):
        try:
            content = ""
            file_size = os.path.getsize(file_path)
            
            if file_path.lower().endswith('.docx'):
                try:
                    import docx2txt
                    content = docx2txt.process(file_path)
                except ImportError:
                    try:
                        from docx import Document as DocxDocument
                        doc = DocxDocument(file_path)
                        content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    except ImportError:
                        print(f"Skipping DOCX {file_path} (install docx2txt or python-docx)")
                        continue
            else:
                # Handle text files with better encoding detection
                encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            content = f.read()
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    print(f"Warning: Could not decode {file_path} with any encoding")
                    continue
            
            # Validate content
            if content and content.strip():
                metadata = {
                    "source": file_path,
                    "file_name": os.path.basename(file_path),
                    "file_type": os.path.splitext(file_path)[1].lower(),
                    "file_size": file_size,
                    "char_count": len(content)
                }
                
                doc = Document(page_content=content, metadata=metadata)
                documents.append(doc)
                
        except Exception as e:
            print(f"Warning: Could not read {file_path}: {e}")
    
    return documents


# format_docs function moved to RAG package


# All node functions moved to RAG package (src/rag/nodes.py)


def main():
    """Main chat loop."""
    parser = argparse.ArgumentParser(description="Interactive RAG chatbot with enhanced memory management")
    parser.add_argument("--docs", default="docs", help="Documents directory")
    parser.add_argument("--thread", default="default", help="Thread ID for persistent memory")
    parser.add_argument("--user", default="default", help="User ID for session management")
    parser.add_argument("--k", type=int, default=8, help="Number of documents to retrieve")
    parser.add_argument("--max-messages", type=int, default=100, help="Maximum messages per thread")
    parser.add_argument("--max-threads", type=int, default=10, help="Maximum threads per user")
    parser.add_argument("--verbose", "-v", action="store_true", help="Enable verbose logging to see workflow state transfers")
    parser.add_argument("--log-level", choices=["DEBUG", "INFO", "WARNING", "ERROR"], default="INFO", help="Set logging level")
    
    # New ingestion parameters
    parser.add_argument("--tenant-id", default="default-tenant", help="Tenant ID for multi-tenancy")
    parser.add_argument("--owner-user-id", default="admin", help="Document owner user ID")
    parser.add_argument("--visibility", choices=["org", "private"], default="org", help="Document visibility scope")
    parser.add_argument("--chunk-size", type=int, default=None, help="Target chunk size (uses config default if not set)")
    parser.add_argument("--chunk-overlap", type=int, default=None, help="Overlap between chunks (uses config default if not set)")
    
    args = parser.parse_args()
    
    # Initialize global memory manager with user settings (will be updated with LLM later)
    global _memory_manager
    _memory_manager = ThreadSafeMemoryManager(max_threads_per_user=args.max_threads)
    
    # Configure logging based on arguments
    if args.verbose:
        log_level = logging.DEBUG
    else:
        log_level = getattr(logging, args.log_level)
    
    # Update logger level
    logger.setLevel(log_level)
    for handler in logger.handlers:
        handler.setLevel(log_level)
    
    # Load environment
    load_dotenv()
    
    # Check for ingestion module
    if not HAS_INGESTION:
        print("ERROR: Ingestion module not found!")
        print("Please install the ingestion module as a submodule or Python package.")
        sys.exit(1)
    
    # Get ingestion configuration
    ingestion_config = get_config()
    
    # Get LLM configuration
    provider = os.getenv("LLM_PROVIDER", "openai").lower()
    model = os.getenv("MODEL", "gpt-4o-mini")
    
    print(f"Using provider: {provider}, model: {model}")
    print(f"Embedding: {ingestion_config.embedding_provider} - {ingestion_config.get_embedding_model()}")
    print(f"Chunking: {ingestion_config.chunk_size} {ingestion_config.chunking_method}s (overlap: {ingestion_config.chunk_overlap})")
    print(f"Memory: max {args.max_messages} messages/thread, {args.max_threads} threads/user")
    if args.chunk_size is not None or args.chunk_overlap is not None:
        print(f"Chunk overrides: size={args.chunk_size}, overlap={args.chunk_overlap}")
    print(f"Tenant: {args.tenant_id}, Owner: {args.owner_user_id}, Visibility: {args.visibility}")
    
    try:
        # Create LLM
        llm = create_llm_provider(provider, model)
        
        # Update memory manager with LLM for summarization
        _memory_manager.llm = llm
        
        # Create MemoryStore for document ingestion and retrieval
        memory_store = MemoryStore()
        
        # Create document context
        doc_ctx = DocumentCtx(
            tenant_id=args.tenant_id,
            owner_user_id=args.owner_user_id,
            document_id="master",  # Will be overridden per document in ingestion
            visibility=args.visibility,
            embedding_version=ingestion_config.get_embedding_model()
        )
        
        # Load and ingest documents into MemoryStore
        total_chunks = load_and_ingest_documents(
            docs_path=args.docs,
            memory_store=memory_store,
            ctx=doc_ctx,
            chunk_size=args.chunk_size,
            overlap=args.chunk_overlap
        )
        
        if total_chunks == 0:
            print("No documents ingested. Please add PDF files to the docs directory.")
            sys.exit(1)
        
        # Create LangChain-compatible retriever using the adapter
        retriever = MemoryRetrieverAdapter(
            store=memory_store,
            ctx=doc_ctx,
            top_k=args.k
        )
        
        # Build RAG workflow
        rag_workflow = create_rag_workflow(llm, retriever, args.k)
        
        # Initialize session variables
        current_user = args.user
        current_thread = args.thread
        
        print(f"\n✅ Ready! User: {current_user}, Thread: {current_thread}")
        print("💬 Enhanced Commands:")
        print("  Basic: /exit, /quit, /reset, /workflow")
        print("  Memory: /threads, /users, /switch, /new, /info")
        print("  Thread: /title <name>, /tag <name>, /cleanup")
        print("  Usage: /switch <thread_id> | /switch user <user_id> [thread_id]")
        if args.verbose:
            print("💡 Verbose logging enabled - you'll see workflow state transfers")
        
        # Set max messages for new threads
        if current_thread in _memory_manager._history_cache:
            _memory_manager._history_cache[current_thread].max_messages = args.max_messages
        
        # Interactive loop
        while True:
            try:
                user_input = input(f"{current_user}> ").strip()
                
                if user_input.lower() in ["/exit", "/quit"]:
                    save_session_history(current_thread)
                    print("Goodbye!")
                    break
                
                if user_input.lower() == "/reset":
                    clear_session_history(current_thread)
                    print("Thread memory cleared.")
                    continue
                
                if user_input.lower() == "/workflow":
                    visualize_graph_structure()
                    continue
                
                # Handle management commands
                handled, new_user, new_thread = handle_management_commands(user_input, current_user, current_thread)
                if handled:
                    # Save current thread before switching
                    save_session_history(current_thread)
                    
                    # Update session variables if changed
                    if new_user:
                        current_user = new_user
                    if new_thread:
                        current_thread = new_thread
                        # Set max messages for new thread
                        history = get_session_history(current_thread, current_user)
                        if hasattr(history, 'max_messages'):
                            history.max_messages = args.max_messages
                    continue
                
                if not user_input:
                    continue
                
                # Process query
                start_time = time.time()
                
                logger.info(f"🚀 Starting RAG workflow: '{user_input}'")
                logger.debug("🚀" * 35)
                logger.debug("🚀 STARTING RAG WORKFLOW")
                logger.debug("🚀" * 35)
                logger.debug(f"User: {current_user}")
                logger.debug(f"Thread: {current_thread}")
                logger.debug(f"Input: '{user_input}'")
                logger.debug(f"Verbose mode: {args.verbose}")
                
                # Get current session history with user context
                session_history = get_session_history(current_thread, current_user)
                logger.debug(f"Session history length: {len(session_history.messages)} messages")
                if hasattr(session_history.metadata, 'summary') and session_history.metadata.summary:
                    logger.debug(f"Has conversation summary: {len(session_history.metadata.summary)} chars")
                
                # Create initial state with user message
                user_message = HumanMessage(content=user_input)
                initial_state = {
                    "messages": session_history.messages + [user_message],
                    "question": user_input,
                    "rewritten_query": "",
                    "context": "",
                    "answer": "",
                    "needs_retrieval": False  # Will be determined by classifier
                }
                
                # Run the workflow
                result = rag_workflow.invoke(initial_state)
                
                # Get the answer
                answer = result["answer"]
                elapsed = time.time() - start_time
                
                logger.info(f"✅ RAG workflow completed in {elapsed:.2f}s")
                logger.debug("🎉" * 35)
                logger.debug("🎉 RAG WORKFLOW COMPLETE")
                logger.debug("🎉" * 35)
                logger.debug(f"Total time: {elapsed:.2f}s")
                logger.debug(f"Final answer length: {len(answer)} chars")
                
                print(f"bot> {answer}")
                print(f"(time: {elapsed:.2f}s)")
                
                # Add messages to session history
                session_history.add_message(user_message)
                session_history.add_message(AIMessage(content=answer))
                
                # Save history after each turn
                save_session_history(current_thread)
                
            except KeyboardInterrupt:
                save_session_history(current_thread)
                print("\nGoodbye!")
                break
            except Exception as e:
                print(f"Error: {e}")
    
    except Exception as e:
        print(f"Startup error: {e}")
        sys.exit(1)


# Optional: Example of how to use the vLLM adapter (commented out)
# from src.adapters.vllm_client_adapter import make_vllm_runnable, SimpleVLLMClient
# 
# def create_vllm_with_adapter():
#     """Alternative vLLM setup using custom adapter."""
#     client = SimpleVLLMClient(
#         base_url=os.getenv("VLLM_BASE_URL"),
#         api_key=os.getenv("VLLM_API_KEY")
#     )
#     return make_vllm_runnable(client)


def visualize_graph_structure():
    """Optional function to visualize the enhanced LangGraph structure."""
    print("\n🔄 Enhanced RAG Workflow Structure (with Conditional Routing):")
    print("┌─────────────────┐")
    print("│   User Query    │")
    print("└─────────┬───────┘")
    print("          │")
    print("   ┌──────▼──────┐")
    print("   │ Classify    │ ← Decides: company docs vs general knowledge")
    print("   │ Query Node  │")
    print("   └──────┬──────┘")
    print("          │")
    print("   ┌──────▼──────┐")
    print("   │ Rewrite     │ ← History-aware rewrite (if needed)")
    print("   │ Query Node  │")
    print("   └──────┬──────┘")
    print("          │")
    print("   ┌──────▼──────┐")
    print("   │ Retrieve    │ ← Search FAISS (if company docs needed)")
    print("   │ Node        │   Otherwise skip")
    print("   └──────┬──────┘")
    print("          │")
    print("   ┌──────▼──────┐")
    print("   │ Generate    │ ← Context-based OR direct answer")
    print("   │ Node        │")
    print("   └──────┬──────┘")
    print("          │")
    print("   ┌──────▼──────┐")
    print("   │   Answer    │")
    print("   └─────────────┘")
    print("\n🎯 Smart routing: Only retrieves when needed!")
    print("💾 Enhanced Memory System:")
    print("   • Thread-safe user/thread management")
    print("   • Auto-cleanup of old conversations") 
    print("   • Configurable memory limits")
    print("   • Rich metadata & tagging support")
    print("   • LangGraph + Enhanced Memory integration")


if __name__ == "__main__":
    main()

